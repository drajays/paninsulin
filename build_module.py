from pathlib import Path
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

src = Path('insulin_education_module.md').read_text(encoding='utf-8')
doc = Document()

# Page layout
for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Theme
styles = doc.styles
styles['Normal'].font.name = 'Arial'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
styles['Normal'].font.size = Pt(9.5)
for name, size, color in [('Title', 25, '0B3B60'), ('Heading 1', 17, '0B3B60'), ('Heading 2', 12, '176B87'), ('Heading 3', 10.5, '333333')]:
    st = styles[name]
    st.font.name = 'Arial'
    st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)

# Add footer page number
for section in doc.sections:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run('Insulin and Diabetes Self-Management Education Module  |  ')
    r.font.size = Pt(8)
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    footer._p.append(fld)

# Helpers
def shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def set_keep(paragraph, value=True):
    pPr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement('w:keepNext')
    keep.set(qn('w:val'), '1' if value else '0')
    pPr.append(keep)

def add_rich(paragraph, text):
    # Basic markdown bold/italic rendering, retaining readable links.
    pos = 0
    pat = re.compile(r'(\*\*.*?\*\*|\*.*?\*)')
    for m in pat.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])

def add_body(text, style=None):
    p = doc.add_paragraph(style=style)
    add_rich(p, text)
    return p

lines = src.splitlines()
# Custom cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_after = Pt(18)
r = p.add_run('INSULIN AND DIABETES\nSELF-MANAGEMENT EDUCATION MODULE')
r.bold = True; r.font.name = 'Arial'; r.font.size = Pt(25); r.font.color.rgb = RGBColor(11,59,96)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('240 patient- and caregiver-focused questions and answers')
r.font.size = Pt(14); r.font.color.rgb = RGBColor(23,107,135)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Evidence review date: July 2026')
r.font.size = Pt(11)

# Cover callout
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(25)
p.paragraph_format.space_after = Pt(10)
add_rich(p, 'For adults, adolescents, families and caregivers. Aligned primarily with ADA Standards of Care in Diabetes—2026, IDF 2025 and relevant international specialty guidance.')
shade(p, 'EAF4F7')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
add_rich(p, 'Important: This is an education resource, not an individualized prescription. Do not start, stop or change insulin, other medicines, pump settings or pregnancy treatment without the treating team. Emergency symptoms such as severe breathing difficulty, confusion, unconsciousness, persistent vomiting, suspected DKA, chest pain or stroke symptoms require urgent care.')
shade(p, 'FFF2CC')

doc.add_page_break()

# Parse content after the initial title, skipping the markdown H1 title line.
# Keep all body material, but turn each numbered chapter into a new page.
first_h1_seen = False
chapter_started = False
for i, line in enumerate(lines):
    if line.startswith('# INSULIN AND DIABETES'):
        first_h1_seen = True
        continue
    if not first_h1_seen:
        continue
    if line.strip() == '':
        continue
    if re.match(r'^# \d+\.', line):
        if chapter_started:
            doc.add_page_break()
        chapter_started = True
        title = line[2:].strip()
        p = doc.add_paragraph(style='Heading 1')
        add_rich(p, title)
        set_keep(p)
        continue
    if line.startswith('## '):
        p = doc.add_paragraph(style='Heading 2')
        add_rich(p, line[3:].strip())
        set_keep(p)
        continue
    if line.startswith('### '):
        p = doc.add_paragraph(style='Heading 3')
        add_rich(p, line[4:].strip())
        set_keep(p)
        continue
    if line.startswith('**') and line.endswith('**'):
        p = doc.add_paragraph()
        add_rich(p, line)
        p.paragraph_format.space_after = Pt(5)
        continue
    if line.startswith('- '):
        p = doc.add_paragraph(style='List Bullet')
        add_rich(p, line[2:])
        continue
    if re.match(r'^\d+\. \*\*', line):
        p = doc.add_paragraph(style='List Number')
        add_rich(p, line)
        continue
    # Answer lines and ordinary paragraphs
    p = doc.add_paragraph()
    add_rich(p, line)
    p.paragraph_format.space_after = Pt(4)

# Set widow/orphan control and spacing for all paragraphs
for p in doc.paragraphs:
    p.paragraph_format.line_spacing = 1.08
    p.paragraph_format.widow_control = True
    for run in p.runs:
        if not run.font.name:
            run.font.name = 'Arial'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

out = Path('insulin_education_module.docx')
doc.save(out)
print(out, out.stat().st_size)
